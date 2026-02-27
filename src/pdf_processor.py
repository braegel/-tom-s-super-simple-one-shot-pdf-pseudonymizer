"""
PDF Processor – Read text from PDFs and redact identified entities.
Uses PyMuPDF (fitz) for all PDF operations.

Key features:
- Uses fitz redaction annotations so the underlying text is truly removed
  (not merely covered by a shape that could be lifted).
- Sorts entities longest-first to avoid partial matches inside longer ones.
- Strips all metadata from the output PDF.
- Enhanced signature / handwriting detection combining five methods:
  1. Embedded image analysis (expanded size heuristics)
  2. Vector drawing clustering (detects pen-drawn signatures)
  3. Ink / freehand annotation detection
  4. PDF form signature field detection
  5. Render-based bottom-zone scan (catch-all for missed elements)
- Logo / brand image detection: repeating images in headers/footers are
  identified and redacted according to the selected mode.
- All PDF metadata is stripped from the output (Author, Creator, Producer, …).
- Multi-format input: accepts PDF, DOCX, DOC, JPG, JPEG.
- Automatic OCR for image-based inputs and PDFs without text layer.
"""

import fitz  # PyMuPDF
from typing import Dict, Tuple, List, Optional, Callable
import os
import json as _json
import re as _re
import subprocess
import tempfile


# ─── Redaction colour palette ───────────────────────────────────────────────
# Pure black redaction boxes – clean, authoritative, no colour distractions.

REDACT_BG       = (0.0, 0.0, 0.0)        # pure black fill
REDACT_FG       = (1.0, 1.0, 1.0)        # white label text


# ---------------------------------------------------------------------------
# Supported input formats
# ---------------------------------------------------------------------------

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".jpg", ".jpeg"}

# Generic legal suffixes – excluded when expanding compound entity names.
_GENERIC_SUFFIXES = {
    "gmbh", "ag", "kg", "ohg", "gbr", "se", "eg", "ev",
    "ltd", "inc", "llc", "co", "ug", "mbh", "e.v.", "e.u.",
    "und", "&", "der", "die", "das", "für", "von",
}


def _has_text_layer(pdf_path: str) -> bool:
    """Check if a PDF has an extractable text layer."""
    doc = fitz.open(pdf_path)
    for page in doc:
        text = page.get_text("text")
        if text.strip():
            doc.close()
            return True
    doc.close()
    return False


def _image_to_pdf(img_path: str) -> str:
    """Convert a JPG/JPEG image file to a single-page PDF.

    Returns the path to a temporary PDF file.
    """
    try:
        tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        tmp.close()
        img_doc = fitz.open(img_path)
        pdf_bytes = img_doc.convert_to_pdf()
        img_doc.close()
        pdf_doc = fitz.open("pdf", pdf_bytes)
        pdf_doc.save(tmp.name)
        pdf_doc.close()
        return tmp.name
    except Exception as e:
        raise RuntimeError(
            f"Fehler beim Konvertieren des Bildes in PDF: {e}\n"
            f"Stellen Sie sicher, dass die Datei ein gültiges JPG/JPEG ist."
        )


def _ocr_pdf(pdf_path: str) -> str:
    """Run OCR on a PDF without a text layer.

    Tries the ``ocrmypdf`` Python API first, then falls back to the CLI.
    Returns the path to the OCR'd temporary PDF.
    """
    tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
    tmp.close()
    try:
        import ocrmypdf
        ocrmypdf.ocr(
            pdf_path, tmp.name,
            language="deu+eng",
            skip_text=True,
            optimize=1,
            progress_bar=False,
        )
        return tmp.name
    except ImportError:
        pass  # try CLI fallback below
    except Exception as e:
        raise RuntimeError(f"OCR-Fehler: {e}")

    # Fallback: command-line ocrmypdf (needs tesseract installed)
    try:
        subprocess.run(
            [
                "ocrmypdf", "--skip-text", "-l", "deu+eng",
                "--optimize", "1", pdf_path, tmp.name,
            ],
            check=True,
            timeout=300,
        )
        return tmp.name
    except FileNotFoundError:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass
        raise RuntimeError(
            "OCR wird benötigt, ist aber nicht installiert.\n"
            "Bitte installieren Sie ocrmypdf und Tesseract:\n"
            "  • pip install ocrmypdf\n"
            "  • Tesseract: apt install tesseract-ocr tesseract-ocr-deu"
        )
    except subprocess.SubprocessError as e:
        raise RuntimeError(f"OCR-Fehler: {e}")


_VISION_OCR_PROMPT = """Extrahiere den VOLLSTÄNDIGEN Text aus diesem Dokument-Scan / Bild.

REGELN:
- Gib den Text EXAKT so wieder wie er im Dokument steht
- Behalte die Struktur bei (Absätze, Aufzählungen, Einrückungen)
- Für Tabellen: verwende | als Spalten-Trenner und neue Zeilen für Reihen
- Überspringe KEINEN Text – auch Fußnoten, Seitenzahlen, Kopfzeilen, Briefköpfe
- Handschriftliche Texte: versuche zu lesen, wenn unleserlich markiere als [HANDSCHRIFT]
- Unterschriften: markiere als [UNTERSCHRIFT]
- Ignoriere reine Grafiken/Bilder/Logos (nur Text extrahieren)
- Gib NUR den extrahierten Text zurück, KEINE Erklärungen oder Kommentare
- Wenn kein Text erkennbar: antworte mit [KEIN TEXT]"""


def _gpt_vision_ocr(pdf_path: str, api_key: str,
                     status_callback: Optional[Callable[[str], None]] = None) -> str:
    """Use GPT-5.2 Vision to extract text from image-based PDF pages.

    Renders each page as a JPEG, sends to GPT-5.2 Vision for text
    extraction.  Much better quality than Tesseract OCR, especially for
    handwriting, complex layouts, and multi-language documents.

    Returns the path to a new text-based PDF containing the extracted text.
    """
    import base64
    from openai import OpenAI

    client = OpenAI(api_key=api_key)
    doc = fitz.open(pdf_path)
    all_text: List[str] = []
    total = len(doc)

    for idx in range(total):
        page = doc[idx]
        if status_callback:
            status_callback(f"KI liest Seite {idx + 1}/{total} …")

        # Render at 200 DPI for good quality
        mat = fitz.Matrix(200.0 / 72.0, 200.0 / 72.0)
        try:
            pix = page.get_pixmap(matrix=mat, colorspace=fitz.csRGB)
        except Exception:
            all_text.append("")
            continue

        img_bytes = pix.tobytes("jpeg")
        b64 = base64.b64encode(img_bytes).decode("utf-8")

        try:
            resp = client.chat.completions.create(
                model="gpt-5.2",
                messages=[{
                    "role": "user",
                    "content": [
                        {"type": "text", "text": _VISION_OCR_PROMPT},
                        {"type": "image_url", "image_url": {
                            "url": f"data:image/jpeg;base64,{b64}",
                            "detail": "high",
                        }},
                    ],
                }],
                temperature=0.0,
                max_completion_tokens=8192,
            )
            page_text = resp.choices[0].message.content.strip()
            if page_text == "[KEIN TEXT]":
                page_text = ""
            all_text.append(page_text)
        except Exception:
            all_text.append("")

    doc.close()
    full_text = "\n\n".join(all_text)

    if not full_text.strip():
        raise ValueError(
            "GPT-5.2 Vision konnte keinen Text im Dokument erkennen."
        )

    return _text_to_pdf(full_text)


def _text_to_pdf(full_text: str) -> str:
    """Create a multi-page PDF from plain text. Returns temp file path."""
    pdf_doc = fitz.open()
    page_w, page_h = 595.28, 841.89
    margin = 50
    text_rect = fitz.Rect(margin, margin, page_w - margin, page_h - margin)

    remaining = full_text
    while remaining.strip():
        page = pdf_doc.new_page(width=page_w, height=page_h)
        rc = page.insert_textbox(
            text_rect, remaining,
            fontname="helv", fontsize=11,
            color=(0, 0, 0),
        )
        if rc >= 0:
            break  # all text fit on this page
        char_capacity = int(len(remaining) * 0.8)
        split_pos = remaining.rfind("\n", 0, char_capacity)
        if split_pos <= 0:
            split_pos = remaining.rfind(" ", 0, char_capacity)
        if split_pos <= 0:
            split_pos = char_capacity
        remaining = remaining[split_pos:].lstrip()

    tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
    tmp.close()
    pdf_doc.save(tmp.name)
    pdf_doc.close()
    return tmp.name


def _extract_docx_text(docx_path: str) -> str:
    """Extract text from a .docx file including paragraphs, tables,
    headers, and footers.  Requires python-docx."""
    from docx import Document as DocxDocument

    doc = DocxDocument(docx_path)
    parts: List[str] = []

    # Headers & footers (all sections)
    for section in doc.sections:
        for hdr in (section.header, section.first_page_header):
            if hdr and hdr.paragraphs:
                for p in hdr.paragraphs:
                    if p.text.strip():
                        parts.append(p.text)
        for ftr in (section.footer, section.first_page_footer):
            if ftr and ftr.paragraphs:
                for p in ftr.paragraphs:
                    if p.text.strip():
                        parts.append(p.text)

    # Main body paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("  |  ".join(cells))

    return "\n\n".join(parts)


def _docx_to_pdf(docx_path: str) -> str:
    """Convert a DOCX/DOC file to PDF.

    Strategy:
      1. Try LibreOffice headless (best quality, preserves formatting)
      2. Fallback: extract text via python-docx and create a simple PDF
         (only works for .docx – .doc requires LibreOffice)
    Returns the path to the resulting PDF.
    """
    ext = os.path.splitext(docx_path)[1].lower()

    # --- Strategy 1: LibreOffice headless ---
    tmp_dir = tempfile.mkdtemp()
    base = os.path.splitext(os.path.basename(docx_path))[0]
    expected = os.path.join(tmp_dir, f"{base}.pdf")
    for lo_cmd in ("libreoffice", "soffice"):
        try:
            subprocess.run(
                [
                    lo_cmd, "--headless", "--convert-to", "pdf",
                    "--outdir", tmp_dir, docx_path,
                ],
                check=True,
                timeout=120,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )
            if os.path.exists(expected):
                return expected
        except (FileNotFoundError, subprocess.SubprocessError):
            continue

    # --- Strategy 2: python-docx text extraction (only .docx) ---
    if ext == ".doc":
        raise RuntimeError(
            "Für .doc-Dateien (altes Word-Format) wird LibreOffice benötigt.\n"
            "Bitte installieren Sie LibreOffice, oder speichern Sie die Datei\n"
            "als .docx und versuchen Sie es erneut."
        )

    try:
        full_text = _extract_docx_text(docx_path)
    except ImportError:
        raise RuntimeError(
            "Für die DOCX-Konvertierung wird LibreOffice oder python-docx benötigt.\n"
            "Bitte installieren Sie eines davon:\n"
            "  • LibreOffice (empfohlen für volle Formatierung)\n"
            "  • pip install python-docx"
        )
    except Exception as e:
        raise RuntimeError(f"Fehler beim Lesen der DOCX-Datei: {e}")

    if not full_text.strip():
        raise ValueError("Das Word-Dokument enthält keinen erkennbaren Text.")

    return _text_to_pdf(full_text)


def _do_ocr(pdf_path: str, api_key: Optional[str],
            status_callback: Optional[Callable[[str], None]] = None) -> str:
    """Extract text from an image-based PDF using GPT Vision (preferred)
    or Tesseract OCR (fallback).  Returns path to a text-based temp PDF."""
    # --- Primary: GPT-5.2 Vision ---
    if api_key:
        try:
            if status_callback:
                status_callback("KI-Texterkennung wird durchgeführt …")
            return _gpt_vision_ocr(pdf_path, api_key, status_callback)
        except Exception:
            pass  # fall through to Tesseract

    # --- Fallback: Tesseract OCR ---
    if status_callback:
        status_callback("OCR wird durchgeführt …")
    return _ocr_pdf(pdf_path)


def prepare_input(
    input_path: str,
    api_key: Optional[str] = None,
    status_callback: Optional[Callable[[str], None]] = None,
) -> str:
    """Prepare an input file for processing.

    Accepts PDF, DOCX, DOC, JPG, and JPEG.
    Converts non-PDF files to PDF.  For image-based content, uses
    GPT-5.2 Vision (if *api_key* given) or Tesseract OCR as fallback
    to extract text.

    Returns the path to a PDF with a text layer.
    If the returned path differs from *input_path*, it is a temporary file
    that the caller should clean up when done.
    """
    ext = os.path.splitext(input_path)[1].lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Nicht unterstütztes Dateiformat: {ext}\n"
            f"Unterstützt: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    if ext in (".jpg", ".jpeg"):
        if status_callback:
            status_callback("Bild wird in PDF konvertiert …")
        pdf_path = _image_to_pdf(input_path)
        result = _do_ocr(pdf_path, api_key, status_callback)
        # Clean up intermediate image PDF
        if result != pdf_path:
            try:
                os.unlink(pdf_path)
            except OSError:
                pass
        return result

    if ext in (".doc", ".docx"):
        if status_callback:
            status_callback("Word-Dokument wird in PDF konvertiert …")
        pdf_path = _docx_to_pdf(input_path)
        # Ensure we have an extractable text layer
        if not _has_text_layer(pdf_path):
            result = _do_ocr(pdf_path, api_key, status_callback)
            if result != pdf_path:
                try:
                    os.unlink(pdf_path)
                except OSError:
                    pass
            return result
        return pdf_path

    # ext == ".pdf"
    if _has_text_layer(input_path):
        return input_path  # already usable

    return _do_ocr(input_path, api_key, status_callback)


# ---------------------------------------------------------------------------
# Scan page detection
# ---------------------------------------------------------------------------

def _page_is_scan(page) -> bool:
    """Detect if *page* is a scanned image (full-page background image).

    On scan pages, the entire visible content is ONE large raster image
    with an invisible OCR text layer on top.  ``apply_redactions`` with
    ``PDF_REDACT_IMAGE_REMOVE`` would delete that image and wipe the
    whole page.  We need ``PDF_REDACT_IMAGE_PIXELS`` instead.
    """
    try:
        images = page.get_images(full=True)
    except Exception:
        return False
    if not images:
        return False

    page_rect = page.rect
    page_area = page_rect.width * page_rect.height
    if page_area < 1:
        return False

    for img_info in images:
        xref = img_info[0]
        try:
            rects = page.get_image_rects(xref)
        except Exception:
            continue
        for rect in rects:
            img_area = rect.width * rect.height
            # Image covers > 50 % of page → almost certainly a scan
            if img_area > page_area * 0.50:
                return True
    return False


def extract_text(pdf_path: str) -> str:
    """Extract the full plain text from a PDF. Requires the PDF to have embedded text."""
    doc = fitz.open(pdf_path)
    pages_text: List[str] = []
    for page in doc:
        pages_text.append(page.get_text("text"))
    doc.close()
    full = "\n".join(pages_text)
    if not full.strip():
        raise ValueError(
            "Das PDF enthält keinen erkennbaren Text. "
            "Bitte stellen Sie sicher, dass das PDF Texterkennung (OCR) hat."
        )
    return full


def _add_redaction(page, rect: fitz.Rect, label: str, mode: str = "pseudo_vars",
                   category: str = ""):
    """Add a redaction annotation to *rect*.

    READABILITY FIRST – the box is kept as tight as possible to the
    original text.  When the label text is longer than the original we
    **shrink the font** (and truncate with "…" as last resort) instead
    of widening the rectangle, so neighbouring text is never covered.

    Returns ``(final_rect, label, font_size, category)`` for the overlay.
    """
    if mode == "anonymize" or not label:
        # fill=REDACT_BG: removes text + blanks scan pixels; overlay adds visuals
        page.add_redact_annot(rect, text="", fill=REDACT_BG)
        return (fitz.Rect(rect), "", 0, category)

    # Target font size: slightly smaller than box height for breathing room.
    # Capped at 9pt for a clean, readable look that doesn't overpower.
    box_h = rect.height
    font_size = min(box_h * 0.68, 9)
    if font_size < 4.0:
        font_size = 4.0

    # Maximum width: original text rect + tiny allowance (4pt).
    # NEVER widen significantly – that would cover adjacent text.
    max_w = rect.width + 4
    padding = 4

    text_w = fitz.get_text_length(label, fontname="helv", fontsize=font_size)

    # Strategy 1: shrink font until it fits (but stay readable)
    while text_w + padding > max_w and font_size > 4.0:
        font_size -= 0.5
        text_w = fitz.get_text_length(label, fontname="helv", fontsize=font_size)

    # Strategy 2: if still too wide, truncate with "…"
    display_label = label
    if text_w + padding > max_w and len(label) > 4:
        while (len(display_label) > 2
               and fitz.get_text_length(display_label + "…", fontname="helv",
                                        fontsize=font_size) + padding > max_w):
            display_label = display_label[:-1]
        display_label = display_label + "…"

    needed_w = fitz.get_text_length(display_label, fontname="helv",
                                    fontsize=font_size) + padding

    # Only extend by the absolute minimum, never more than 4pt extra
    final_rect = fitz.Rect(rect)
    if needed_w > rect.width:
        extra = min(needed_w - rect.width, 4)
        page_rect = page.rect
        new_x1 = min(rect.x0 + rect.width + extra, page_rect.width - 2)
        final_rect = fitz.Rect(rect.x0, rect.y0, new_x1, rect.y1)

    # fill=REDACT_BG: removes text + blanks scan pixels; overlay adds visuals
    page.add_redact_annot(
        final_rect,
        text="",
        fill=REDACT_BG,
    )
    return (fitz.Rect(final_rect), display_label, font_size, category)


# ---------------------------------------------------------------------------
# Signature / Handwriting Detection (enhanced)
# ---------------------------------------------------------------------------

# Bottom fraction of each page treated as the "signature zone" where
# detection is more aggressive (signatures, initials, stamps).
_SIG_ZONE_FRACTION = 0.40

# Top fraction of each page – letterheads sometimes contain initials / stamps.
_HEADER_ZONE_FRACTION = 0.12

# Maximum gap (pt) between drawing strokes to consider them one cluster.
_CLUSTER_GAP = 20

# Minimum number of vector strokes in a cluster to flag as handwriting.
_MIN_CLUSTER_STROKES = 2

# Extra padding (pt) around every detected signature element.
_REDACT_MARGIN = 2


def _expand_rect(rect: fitz.Rect, page_rect: fitz.Rect, margin: float = 15) -> fitz.Rect:
    """Expand *rect* by *margin* in all directions, clamped to page bounds."""
    return fitz.Rect(
        max(page_rect.x0, rect.x0 - margin),
        max(page_rect.y0, rect.y0 - margin),
        min(page_rect.x1, rect.x1 + margin),
        min(page_rect.y1, rect.y1 + margin),
    )


def _safe_expand_rect(rect: fitz.Rect, page, margin: float = 2) -> fitz.Rect:
    """Expand *rect* by *margin* but shrink back if it would overlap text
    blocks that lie OUTSIDE the original rect.

    This guarantees that redaction boxes never obscure adjacent readable
    text – the single most important rule for output readability.
    """
    page_rect = page.rect
    expanded = _expand_rect(rect, page_rect, margin)

    try:
        blocks = page.get_text("blocks")
    except Exception:
        return expanded

    for block in blocks:
        br = fitz.Rect(block[:4])
        if br.is_empty or not expanded.intersects(br):
            continue
        # Text block already inside the original rect → target text, OK
        if rect.contains(br):
            continue
        # Block partially overlaps the expansion – clamp each edge
        if br.x1 > rect.x0 and br.x0 < rect.x0 and expanded.x0 < br.x1:
            expanded = fitz.Rect(max(expanded.x0, br.x1 + 0.5),
                                 expanded.y0, expanded.x1, expanded.y1)
        if br.x0 < rect.x1 and br.x1 > rect.x1 and expanded.x1 > br.x0:
            expanded = fitz.Rect(expanded.x0, expanded.y0,
                                 min(expanded.x1, br.x0 - 0.5), expanded.y1)
        if br.y1 > rect.y0 and br.y0 < rect.y0 and expanded.y0 < br.y1:
            expanded = fitz.Rect(expanded.x0, max(expanded.y0, br.y1 + 0.5),
                                 expanded.x1, expanded.y1)
        if br.y0 < rect.y1 and br.y1 > rect.y1 and expanded.y1 > br.y0:
            expanded = fitz.Rect(expanded.x0, expanded.y0,
                                 expanded.x1, min(expanded.y1, br.y0 - 0.5))

    # If clamping collapsed the rect, fall back to original
    if expanded.is_empty or expanded.width < 2 or expanded.height < 2:
        return fitz.Rect(rect)
    return expanded


def _cluster_rects(rects: List[fitz.Rect], max_gap: float = 20):
    """Group overlapping / nearby rectangles into clusters.

    Returns a list of ``(merged_rect, member_count)`` tuples.
    """
    if not rects:
        return []

    clusters: List[List] = [[fitz.Rect(rects[0]), 1]]

    for rect in rects[1:]:
        merged = False
        for cluster in clusters:
            padded = fitz.Rect(
                cluster[0].x0 - max_gap, cluster[0].y0 - max_gap,
                cluster[0].x1 + max_gap, cluster[0].y1 + max_gap,
            )
            if padded.intersects(rect):
                cluster[0] |= rect
                cluster[1] += 1
                merged = True
                break
        if not merged:
            clusters.append([fitz.Rect(rect), 1])

    # Iteratively merge clusters that now overlap each other.
    changed = True
    while changed:
        changed = False
        new_clusters: List[List] = []
        used: set = set()
        for i, (r1, c1) in enumerate(clusters):
            if i in used:
                continue
            for j in range(i + 1, len(clusters)):
                if j in used:
                    continue
                r2, c2 = clusters[j]
                padded = fitz.Rect(
                    r1.x0 - max_gap, r1.y0 - max_gap,
                    r1.x1 + max_gap, r1.y1 + max_gap,
                )
                if padded.intersects(r2):
                    r1 |= r2
                    c1 += c2
                    used.add(j)
                    changed = True
            new_clusters.append([r1, c1])
            used.add(i)
        clusters = new_clusters

    return [(c[0], c[1]) for c in clusters]


# -- Individual detection methods ------------------------------------------

def _image_looks_like_signature(xref: int, doc) -> bool:
    """Heuristic: check if an image has low colour diversity (typical for
    handwriting / stamps which are mostly one ink colour on white/transparent).
    Returns True if the image is *likely* a signature rather than a photo.

    Uses pixel sampling (max 4000 pixels) for speed on large images.
    Thresholds are intentionally loose.
    """
    try:
        pix = fitz.Pixmap(doc, xref)
        if pix.n > 1:
            pix = fitz.Pixmap(fitz.csGRAY, pix)
        samples = pix.samples
        total = len(samples)
        if total < 10:
            return False
        # Sample up to 4000 pixels evenly for speed
        if total > 4000:
            step = total // 4000
            sampled = samples[::step]
            n = len(sampled)
        else:
            sampled = samples
            n = total
        dark = sum(1 for b in sampled if b < 100)
        light = sum(1 for b in sampled if b > 180)
        ratio = (dark + light) / n
        if ratio > 0.70:
            return True
        dark_ratio = dark / n
        light_ratio = light / n
        return dark_ratio > 0.03 and light_ratio > 0.60
    except Exception:
        return False


def _redact_signature_images(page):
    """Detect and redact images that look like signatures or handwriting.

    Extra aggressive in the signature zone (bottom 40%) and header zone
    (top 12%).  Uses pixel analysis to distinguish signatures from photos.
    Intentionally generous – a missed signature is worse than a false positive.
    """
    page_rect = page.rect
    sig_zone_top = page_rect.height * (1 - _SIG_ZONE_FRACTION)
    header_zone_bottom = page_rect.height * _HEADER_ZONE_FRACTION

    try:
        images = page.get_images(full=True)
    except Exception:
        return

    doc = page.parent

    for img_info in images:
        xref = img_info[0]
        try:
            rects = page.get_image_rects(xref)
        except Exception:
            continue
        for rect in rects:
            w, h = rect.width, rect.height
            area = w * h
            in_sig_zone = rect.y0 >= sig_zone_top
            in_header_zone = rect.y1 <= header_zone_bottom

            # Skip very large images (full-page scans, large photos)
            if w > page_rect.width * 0.75 and h > page_rect.height * 0.4:
                continue

            # General signature detection (anywhere on page):
            # Generous size range – signatures vary enormously in size.
            if 10 < w < 500 and 4 < h < 200 and 100 < area < 120_000:
                if _image_looks_like_signature(xref, doc):
                    expanded = _safe_expand_rect(rect, page, _REDACT_MARGIN)
                    page.add_redact_annot(expanded, text="", fill=REDACT_BG)
                    continue

            # In signature/header zone: even more generous for small marks
            # (initials, paraphs, stamps, small scribbles)
            if (in_sig_zone or in_header_zone) and w > 8 and h > 4 and area > 80:
                if w < page_rect.width * 0.6 and h < page_rect.height * 0.20:
                    if _image_looks_like_signature(xref, doc):
                        expanded = _safe_expand_rect(rect, page, _REDACT_MARGIN)
                        page.add_redact_annot(expanded, text="", fill=REDACT_BG)


def _redact_signature_drawings(page):
    """Detect clusters of vector paths that resemble handwriting.

    Hand-drawn signatures consist of many short, often curved strokes
    clustered together – very different from table borders or form lines.
    Checks the entire page, with lower thresholds in signature/header zones.
    """
    try:
        drawings = page.get_drawings()
    except Exception:
        return

    if not drawings:
        return

    page_rect = page.rect
    sig_zone_top = page_rect.height * (1 - _SIG_ZONE_FRACTION)
    header_zone_bottom = page_rect.height * _HEADER_ZONE_FRACTION

    stroke_rects: List[fitz.Rect] = []
    for d in drawings:
        rect = fitz.Rect(d["rect"])

        # Skip page-wide horizontal rules / table borders
        if rect.width > page_rect.width * 0.6 and rect.height < 3:
            continue
        # Skip full-height elements (side borders)
        if rect.height > page_rect.height * 0.35:
            continue
        # Skip large filled backgrounds / boxes
        if rect.width * rect.height > page_rect.width * page_rect.height * 0.20:
            continue
        # Skip invisible specks
        if rect.width < 0.5 and rect.height < 0.5:
            continue

        stroke_rects.append(rect)

    if not stroke_rects:
        return

    clusters = _cluster_rects(stroke_rects, max_gap=_CLUSTER_GAP)

    for cluster_rect, stroke_count in clusters:
        in_hotzone = (cluster_rect.y0 >= sig_zone_top
                      or cluster_rect.y1 <= header_zone_bottom)
        # In signature/header zone: very sensitive (2 strokes).
        # Elsewhere: slightly higher threshold to avoid decorative elements.
        min_strokes = _MIN_CLUSTER_STROKES if in_hotzone else _MIN_CLUSTER_STROKES + 1

        if stroke_count < min_strokes:
            continue

        # Skip overly large clusters that are probably layout elements
        if cluster_rect.width > page_rect.width * 0.5 and cluster_rect.height > 120:
            continue

        expanded = _safe_expand_rect(cluster_rect, page, _REDACT_MARGIN)
        page.add_redact_annot(expanded, text="", fill=REDACT_BG)


def _redact_ink_annotations(page):
    """Detect and redact Ink (freehand) annotations – digital pen signatures."""
    try:
        annot = page.first_annot
    except Exception:
        return

    while annot:
        try:
            # PDF annotation type 19 = Ink (freehand drawing)
            if annot.type[0] == 19:
                expanded = _safe_expand_rect(fitz.Rect(annot.rect), page, _REDACT_MARGIN)
                page.add_redact_annot(expanded, text="", fill=REDACT_BG)
            annot = annot.next
        except Exception:
            break


def _redact_form_signature_fields(page):
    """Detect and redact PDF form signature widgets."""
    try:
        widget = page.first_widget
    except Exception:
        return

    while widget:
        try:
            # field_type 7 = signature field in PyMuPDF
            if widget.field_type == 7:
                expanded = _safe_expand_rect(fitz.Rect(widget.rect), page, _REDACT_MARGIN)
                page.add_redact_annot(expanded, text="", fill=REDACT_BG)
            widget = widget.next
        except Exception:
            break


def _redact_bottom_zone_scan(page):
    """Render-based catch-all: render the signature zone at low resolution
    and look for dark marks in areas that contain no text blocks.

    This catches signatures embedded as XObjects, unusual image formats,
    or any other visual element the targeted methods above may miss.
    """
    page_rect = page.rect
    sig_zone_top = page_rect.height * (1 - _SIG_ZONE_FRACTION)

    # Collect text-block rectangles inside the signature zone
    text_rects: List[fitz.Rect] = []
    try:
        for block in page.get_text("blocks"):
            r = fitz.Rect(block[:4])
            if r.y1 > sig_zone_top:
                text_rects.append(r)
    except Exception:
        pass

    # Render at 72 DPI – sufficient for dark-mark detection, fast
    scale = 1.0
    mat = fitz.Matrix(scale, scale)
    clip = fitz.Rect(page_rect.x0, sig_zone_top, page_rect.x1, page_rect.y1)

    try:
        pix = page.get_pixmap(matrix=mat, clip=clip, colorspace=fitz.csGRAY)
    except Exception:
        return

    pw, ph = pix.width, pix.height
    if pw < 2 or ph < 2:
        return

    samples = pix.samples  # grayscale bytes (one byte per pixel)

    # Divide into a grid of ~30px cells (larger = faster, still good detection)
    n_cols = max(1, pw // 30)
    n_rows = max(1, ph // 30)
    cell_w = pw / n_cols
    cell_h = ph / n_rows

    suspect_cells: List[fitz.Rect] = []

    for row in range(n_rows):
        for col in range(n_cols):
            # Map cell back to page coordinates
            cell_x0 = page_rect.x0 + col * cell_w / scale
            cell_y0 = sig_zone_top + row * cell_h / scale
            cell_x1 = cell_x0 + cell_w / scale
            cell_y1 = cell_y0 + cell_h / scale
            cell_page_rect = fitz.Rect(cell_x0, cell_y0, cell_x1, cell_y1)

            # Skip cells that overlap known text blocks
            if any(cell_page_rect.intersects(tr) for tr in text_rects):
                continue

            # Count dark pixels in this cell
            px0 = int(col * cell_w)
            py0 = int(row * cell_h)
            px1 = min(int((col + 1) * cell_w), pw)
            py1 = min(int((row + 1) * cell_h), ph)

            dark = 0
            total = 0
            for y in range(py0, py1):
                offset = y * pw
                for x in range(px0, px1):
                    if samples[offset + x] < 120:
                        dark += 1
                    total += 1

            # Flag cells where > 10 % of pixels are dark (non-text marks)
            # Raised from 5% to reduce false positives on light backgrounds
            if total > 0 and dark / total > 0.10:
                suspect_cells.append(cell_page_rect)

    if not suspect_cells:
        return

    # Merge adjacent suspect cells and redact broad areas
    clusters = _cluster_rects(suspect_cells, max_gap=14)
    for merged_rect, count in clusters:
        if count >= 3:
            expanded = _safe_expand_rect(merged_rect, page, _REDACT_MARGIN)
            page.add_redact_annot(expanded, text="", fill=REDACT_BG)


# ---------------------------------------------------------------------------
# Logo / brand image / letterhead detection
# ---------------------------------------------------------------------------

_FOOTER_ZONE_FRACTION = 0.12   # bottom 12 %


def _find_repeating_image_xrefs(doc) -> set:
    """Pre-scan all pages to find image xrefs that appear on 2+ pages.

    Images repeated across pages are very likely logos / letterheads.
    """
    xref_page_count: dict = {}
    for page_idx in range(len(doc)):
        page = doc[page_idx]
        try:
            images = page.get_images(full=True)
        except Exception:
            continue
        seen_on_page: set = set()
        for img_info in images:
            xref = img_info[0]
            if xref not in seen_on_page:
                seen_on_page.add(xref)
                xref_page_count[xref] = xref_page_count.get(xref, 0) + 1
    return {xref for xref, count in xref_page_count.items() if count >= 2}


def _redact_logo_images(page, repeating_xrefs: set, mode: str) -> int:
    """Detect and redact logo / brand images and letterheads on *page*.

    Aggressively targets:
      • ANY image in the header zone  (top 20 %) – letterheads, logos, brand graphics
      • Repeating images anywhere – branding that appears on multiple pages
      • Small images in the footer zone  (bottom 12 %)

    Only preserves large content images in the body area.
    Returns the number of logo redactions added.
    """
    page_rect = page.rect
    header_bottom = page_rect.y0 + page_rect.height * _HEADER_ZONE_FRACTION
    footer_top = page_rect.y0 + page_rect.height * (1 - _FOOTER_ZONE_FRACTION)

    try:
        images = page.get_images(full=True)
    except Exception:
        return 0

    count = 0
    for img_info in images:
        xref = img_info[0]
        try:
            rects = page.get_image_rects(xref)
        except Exception:
            continue

        for rect in rects:
            w, h = rect.width, rect.height

            # Skip full-page backgrounds
            if w > page_rect.width * 0.9 and h > page_rect.height * 0.5:
                continue
            # Skip tiny invisible elements
            if w < 3 or h < 3:
                continue

            is_repeating = xref in repeating_xrefs
            in_header = rect.y1 <= header_bottom
            in_footer = rect.y0 >= footer_top

            should_redact = False

            if in_header:
                # Only auto-redact repeating brand images or small logo-
                # sized images in the header zone – large content images
                # (photos, infographics) are left alone.
                if is_repeating or (w < 250 and h < 120):
                    should_redact = True
            elif is_repeating:
                # Repeating image anywhere = branding
                should_redact = True
            elif in_footer and h < 100:
                # Images in footer zone = likely footer branding
                should_redact = True

            if should_redact:
                expanded = _expand_rect(rect, page_rect, 2)
                # Subtle fill, no "LOGO" text – keeps output clean and subtle.
                # fill=REDACT_BG: annotation removes content, overlay handles visuals.
                page.add_redact_annot(expanded, text="", fill=REDACT_BG)
                count += 1

    return count


def _redact_header_zone_drawings(page):
    """Redact vector drawings in the header zone (letterhead graphics, lines, etc.).

    Catches non-image letterhead elements like decorative lines, shapes,
    and vector logos that are not embedded as raster images.
    """
    page_rect = page.rect
    header_bottom = page_rect.y0 + page_rect.height * _HEADER_ZONE_FRACTION

    try:
        drawings = page.get_drawings()
    except Exception:
        return

    if not drawings:
        return

    header_strokes: List[fitz.Rect] = []
    for d in drawings:
        rect = fitz.Rect(d["rect"])
        # Only consider drawings in the header zone
        if rect.y1 > header_bottom:
            continue
        # Skip invisible specks
        if rect.width < 2 and rect.height < 2:
            continue
        # Skip page-wide single hairlines (likely just a separator)
        if rect.width > page_rect.width * 0.7 and rect.height < 2:
            continue
        header_strokes.append(rect)

    if not header_strokes:
        return

    # Cluster nearby drawings and redact clusters with multiple strokes
    clusters = _cluster_rects(header_strokes, max_gap=8)
    for cluster_rect, stroke_count in clusters:
        if stroke_count >= 2:
            expanded = _expand_rect(cluster_rect, page_rect, 2)
            page.add_redact_annot(expanded, text="", fill=REDACT_BG)


# ---------------------------------------------------------------------------
# Metadata stripping
# ---------------------------------------------------------------------------

def _strip_metadata(doc):
    """Remove all identifying metadata from the PDF document."""
    doc.set_metadata({
        "producer": "",
        "format": "",
        "encryption": "",
        "author": "",
        "modDate": "",
        "keywords": "",
        "title": "",
        "creationDate": "",
        "creator": "",
        "subject": "",
        "trapped": "",
    })
    try:
        doc.del_xml_metadata()
    except Exception:
        pass


# -- Orchestrator -----------------------------------------------------------

def _detect_and_redact_signatures(page, is_scan: bool = False):
    """Run signature / handwriting detection methods on *page*.

    On scan pages, image analysis, vector clustering, and the render-based
    bottom-zone scan are SKIPPED because they would analyse the scan
    image itself and produce massive false positives.  Only ink
    annotations, form fields, and (externally) vision-based detection
    are used for scans.
    """
    if not is_scan:
        _redact_signature_images(page)
        _redact_signature_drawings(page)
    _redact_ink_annotations(page)
    _redact_form_signature_fields(page)
    if not is_scan:
        _redact_bottom_zone_scan(page)


# ---------------------------------------------------------------------------
# Smart vision page selection – avoid unnecessary API calls
# ---------------------------------------------------------------------------

# Keywords in page text that suggest a signature / handwriting may be present.
_SIG_HINT_WORDS = _re.compile(
    r"unterschrift|signatur|gez\.|gezeichnet|unterzeichn|handzeichen|"
    r"vollmacht|bevollmächtigt|hiermit bestätig|eigenhändig|"
    r"signature|signed|witness",
    _re.IGNORECASE,
)


def _page_needs_vision(page, page_idx: int, total_pages: int) -> bool:
    """Decide whether a page is worth sending to the Vision API.

    Returns True for:
    - First and last page (almost always contain signatures/logos)
    - Second page and second-to-last (common for multi-page contracts)
    - Any page whose text contains signature-related keywords
    - Documents with 4 or fewer pages (scan everything)
    """
    if total_pages <= 4:
        return True
    if page_idx <= 1 or page_idx >= total_pages - 2:
        return True
    try:
        text = page.get_text("text")
        if _SIG_HINT_WORDS.search(text):
            return True
    except Exception:
        pass
    return False


# ---------------------------------------------------------------------------
# GPT-5.2 Vision-based signature detection  (catch-all for missed handwriting)
# ---------------------------------------------------------------------------

_VISION_PROMPT = """Du bist ein Experte für Dokumenten-Anonymisierung. Analysiere dieses Dokumentbild SEHR GRÜNDLICH auf handschriftliche und visuelle Elemente, die Personen identifizieren.

SUCHE INTENSIV NACH (jeweils mit Typ angeben):

1. type="unterschrift" – JEDE handschriftliche Unterschrift, Signatur, Namenszug.
   Das sind geschwungene Tintenstriche, typische Unterschriften-Kurven, auch:
   - Flüchtige, kaum lesbare Unterschriften (einzelner Strich mit Schnörkel)
   - Sehr kleine Unterschriften in Ecken oder am Seitenrand
   - Unterschriften die teilweise von Linien oder Text überlagert werden
   - Unterschriften auf dunklem oder farbigem Hintergrund
   - Blaue, schwarze, oder andere Tintenfarben
   - Digitale Unterschriften die handschriftlich aussehen

2. type="paraphe" – Handschriftliche Kürzel, Initialen, Häkchen, Paraphen.
   JEDE handschriftliche Markierung: Haken, Kreuze, Kürzel, "gez.", Initialen
   wie "HM" oder "S.W.", Randnotizen in Handschrift, auch wenn winzig klein.

3. type="logo" – Firmenlogos, Markenzeichen, Wappen, graphische Briefkopf-
   Elemente (NICHT rein typographische Firmennamen, nur echte Grafiken/Icons)

4. type="stempel" – Firmenstempel, Amtsstempel, Siegel, Rundstempel, Datumsstempel.
   Auch verblasste oder teilweise sichtbare Stempel.

5. type="foto" – Passfotos, Profilbilder, eingescannte Fotos von Personen

NICHT MELDEN:
- Gedruckten maschinellen Text (auch kursiv oder fett)
- Tabellen-Linien, Rahmen, Trennstriche, Formular-Linien
- Seitenzahlen, Fußnoten
- Große Hintergrundmuster, ganzseitige Wasserzeichen

OBERSTE REGEL: Im Zweifel IMMER melden! Eine übersehene Unterschrift ist
ein schwerer Datenschutzverstoß. Ein falscher Treffer kann korrigiert werden.
Prüfe JEDE Ecke, JEDEN Rand, JEDE Zeile wo "Unterschrift:" oder ähnliches steht.

Gib für jede Fundstelle eine passende Bounding-Box (Prozent der Seitenbreite/-höhe).

Antworte NUR mit JSON:
{
  "signatures": [
    {"x_pct": 10, "y_pct": 85, "w_pct": 25, "h_pct": 5, "type": "unterschrift"}
  ]
}

Wenn NICHTS gefunden wird: {"signatures": []}"""


def _detect_visuals_with_vision(page, api_key: str) -> List[Tuple[fitz.Rect, str]]:
    """Use GPT-5.2 vision to detect signatures, logos, stamps on *page*.

    Renders the page as a JPEG, sends it to the vision model, and
    returns a list of ``(fitz.Rect, type_str)`` tuples.
    *type_str* is one of: unterschrift, paraphe, logo, stempel, foto.
    """
    import base64

    page_rect = page.rect

    # Render at 200 DPI – good detail for signatures, fast upload (vs 300 DPI)
    scale = 200.0 / 72.0
    mat = fitz.Matrix(scale, scale)
    try:
        pix = page.get_pixmap(matrix=mat, colorspace=fitz.csRGB)
    except Exception:
        return []

    # Convert to JPEG (quality 80 – good balance of detail and speed)
    img_bytes = pix.tobytes("jpeg", jpg_quality=80)
    b64_image = base64.b64encode(img_bytes).decode("utf-8")

    # Call GPT-5.2 vision
    try:
        from openai import OpenAI
        client = OpenAI(api_key=api_key)
        response = client.chat.completions.create(
            model="gpt-5.2",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": _VISION_PROMPT},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{b64_image}",
                                "detail": "high",
                            },
                        },
                    ],
                },
            ],
            temperature=0.0,
            max_completion_tokens=2048,
        )
    except Exception:
        return []

    # Parse response
    try:
        text = response.choices[0].message.content.strip()
        fence = _re.search(r"```(?:json)?\s*\n?(.*?)```", text, _re.DOTALL)
        if fence:
            text = fence.group(1).strip()
        data = _json.loads(text)
        sigs = data.get("signatures", [])
    except Exception:
        return []

    # Convert percentage-based bboxes to page coordinates
    results: List[Tuple[fitz.Rect, str]] = []
    pw = page_rect.width
    ph = page_rect.height
    for sig in sigs:
        try:
            x = page_rect.x0 + sig["x_pct"] / 100.0 * pw
            y = page_rect.y0 + sig["y_pct"] / 100.0 * ph
            w = sig["w_pct"] / 100.0 * pw
            h = sig["h_pct"] / 100.0 * ph
            rect = fitz.Rect(x, y, x + w, y + h)
            sig_type = sig.get("type", "unterschrift")
            # Sanity check: not the entire page; allow very small marks (paraphs)
            if rect.width > 3 and rect.height > 2 and rect.width < pw * 0.95:
                results.append((rect, sig_type))
        except (KeyError, TypeError):
            continue

    return results


_KAPPA = 0.5522847498  # cubic Bézier approximation of a quarter-circle


def _draw_rounded_rect(shape, rect: fitz.Rect, radius: float = 2.0):
    """Draw a rectangle with slightly rounded corners into *shape*.

    Uses cubic Bézier arcs at each corner for a smooth, professional look.
    *radius* is clamped so it never exceeds half the box dimensions.
    """
    r = min(radius, rect.width / 2, rect.height / 2, 3.0)
    if r < 0.5:
        shape.draw_rect(rect)
        return

    x0, y0, x1, y1 = rect.x0, rect.y0, rect.x1, rect.y1
    k = r * _KAPPA

    # Build path clockwise from top-left corner
    # Top edge
    shape.draw_line(fitz.Point(x0 + r, y0), fitz.Point(x1 - r, y0))
    # Top-right corner
    shape.draw_bezier(
        fitz.Point(x1 - r, y0),
        fitz.Point(x1 - r + k, y0), fitz.Point(x1, y0 + r - k),
        fitz.Point(x1, y0 + r))
    # Right edge
    shape.draw_line(fitz.Point(x1, y0 + r), fitz.Point(x1, y1 - r))
    # Bottom-right corner
    shape.draw_bezier(
        fitz.Point(x1, y1 - r),
        fitz.Point(x1, y1 - r + k), fitz.Point(x1 - r + k, y1),
        fitz.Point(x1 - r, y1))
    # Bottom edge
    shape.draw_line(fitz.Point(x1 - r, y1), fitz.Point(x0 + r, y1))
    # Bottom-left corner
    shape.draw_bezier(
        fitz.Point(x0 + r, y1),
        fitz.Point(x0 + r - k, y1), fitz.Point(x0, y1 - r + k),
        fitz.Point(x0, y1 - r))
    # Left edge
    shape.draw_line(fitz.Point(x0, y1 - r), fitz.Point(x0, y0 + r))
    # Top-left corner
    shape.draw_bezier(
        fitz.Point(x0, y0 + r),
        fitz.Point(x0, y0 + r - k), fitz.Point(x0 + r - k, y0),
        fitz.Point(x0 + r, y0))


def _draw_redaction_overlays(page, overlays: list):
    """Draw clean black redaction boxes over areas where content was removed.

    Annotations use ``fill=REDACT_BG`` which (a) deletes text content and
    (b) blanks scan-image pixels.  This overlay step adds the polished
    visual treatment on top, guaranteeing consistent appearance on every
    PDF type – both native text PDFs and scanned documents.

    Design: pure black fill, slightly rounded corners, centred white label.
    """
    if not overlays:
        return

    for rect, label, font_size, _category in overlays:
        # ── Filled background with rounded corners ──
        shape = page.new_shape()
        _draw_rounded_rect(shape, rect, radius=2.0)
        shape.finish(fill=REDACT_BG, color=REDACT_BG, width=0)
        shape.commit()

        # ── Text label (pseudo modes) ──
        if label and font_size > 0:
            text_w = fitz.get_text_length(label, fontname="helv", fontsize=font_size)
            text_x = rect.x0 + max(0, (rect.width - text_w) / 2)
            # Precise vertical centre: baseline = centre + ascender/2
            text_y = rect.y0 + (rect.height + font_size * 0.72) / 2
            page.insert_text(
                fitz.Point(text_x, text_y),
                label,
                fontname="helv",
                fontsize=font_size,
                color=REDACT_FG,
            )


# ---------------------------------------------------------------------------
# Legal reference / numbering protection
# ---------------------------------------------------------------------------

# Patterns that must NEVER be redacted – §§, article numbers, outline
# numbering like "1.1.", "aa)", "1.c)", "III.", roman numerals, etc.
_LEGAL_PROTECT_PATTERNS = [
    _re.compile(r"^§+\s*\d"),                      # § 1, §§ 12
    _re.compile(r"^Art\.?\s*\d"),                   # Art. 5, Art 12
    _re.compile(r"^\d{1,4}\.\d"),                   # 1.1, 12.3.4
    _re.compile(r"^\d{1,4}\.$"),                    # 1. 2. 3.
    _re.compile(r"^\d{1,4}\)$"),                    # 1) 2) 3)
    _re.compile(r"^[a-z]{1,3}\)$"),                 # a) aa) ab)
    _re.compile(r"^[a-z]{1,3}\.$"),                 # a. b. c.
    _re.compile(r"^\d+\.[a-z]\)"),                  # 1.a) 2.c)
    _re.compile(r"^[IVXLCDM]{1,6}\.?$"),           # I. II. IV. XII
    _re.compile(r"^Abs\.?\s*\d"),                   # Abs. 1, Abs 3
    _re.compile(r"^Nr\.?\s*\d"),                    # Nr. 1, Nr 3
    _re.compile(r"^Ziff\.?\s*\d"),                  # Ziff. 1
    _re.compile(r"^lit\.?\s*[a-z]"),                # lit. a, lit b
    _re.compile(r"^\(\d{1,3}\)$"),                  # (1) (2) (12)
    _re.compile(r"^\([a-z]{1,3}\)$"),               # (a) (aa) (ab)
]


def _is_legal_numbering(text: str) -> bool:
    """Return True if *text* looks like a legal reference or outline number
    that must never be redacted."""
    t = text.strip()
    if not t:
        return False
    return any(p.match(t) for p in _LEGAL_PROTECT_PATTERNS)


def _expand_entity_map(entity_map: Dict[str, Tuple[str, str]]):
    """Derive additional PII fragments from compound entities.

    Mutates *entity_map* in place.  Generates:
    - Sub-word fragments (individual significant words from multi-word entities)
    - Name order variants ("Hans Müller" → "Müller, Hans", "Müller Hans")
    - Initial forms ("Hans Müller" → "H. Müller", "H.Müller")
    - Case variants (UPPER / lower)
    - Spaced / gesperrt variants ("S P A R K A S S E")

    The goal is to catch the same PII even when it appears in a shorter,
    reordered, abbreviated, or differently cased form elsewhere in the document.
    """
    additions: Dict[str, Tuple[str, str]] = {}
    existing = set(entity_map.keys())

    # ── Titles / salutations to strip when generating sub-words ──
    _NAME_TITLES = {"dr", "prof", "herr", "frau", "mr", "mrs", "ms",
                    "von", "van", "de", "zu", "di", "dr.", "prof."}

    for text, (label, cat) in list(entity_map.items()):
        words = text.split()

        if cat == "UNTERNEHMEN" and len(words) >= 2:
            # For compound company names, add significant sub-words.
            for w in words:
                wl = w.lower().rstrip(".,;:")
                if len(wl) >= 4 and wl not in _GENERIC_SUFFIXES and w not in existing:
                    if not _is_legal_numbering(w):
                        additions[w] = (label, cat)

        elif cat in ("VORNAME", "NACHNAME") and len(words) >= 2:
            # ── Sub-word extraction ──
            # "Dr. Hans Müller" → "Hans", "Müller"
            name_parts = []
            for w in words:
                wl = w.rstrip(".")
                if (len(wl) >= 2 and w not in existing
                        and wl.lower() not in _NAME_TITLES):
                    if not _is_legal_numbering(w):
                        additions[w] = (label, cat)
                        name_parts.append(w)
                elif wl.lower() not in _NAME_TITLES and len(wl) >= 2:
                    name_parts.append(w)

            # ── Name order variants ──
            # "Hans Müller" → "Müller, Hans" and "Müller Hans"
            pure_parts = [w for w in words if w.rstrip(".").lower() not in _NAME_TITLES]
            if len(pure_parts) == 2:
                a, b = pure_parts
                reversed_comma = f"{b}, {a}"
                reversed_plain = f"{b} {a}"
                for var in (reversed_comma, reversed_plain):
                    if var != text and var not in existing and var not in additions:
                        additions[var] = (label, cat)

                # ── Initial forms ──
                # "Hans Müller" → "H. Müller", "H.Müller"
                initial_a = a[0] + "."
                initial_b = b[0] + "."
                for var in (f"{initial_a} {b}", f"{initial_a}{b}",
                            f"{a} {initial_b}", f"{initial_b} {a}",
                            f"{initial_a} {initial_b}"):
                    if len(var) >= 3 and var not in existing and var not in additions:
                        additions[var] = (label, cat)

            elif len(pure_parts) == 3:
                # "Hans Peter Müller" → "H. P. Müller", "H.P. Müller"
                a, b, c = pure_parts
                for var in (f"{a[0]}. {b[0]}. {c}", f"{a[0]}.{b[0]}. {c}",
                            f"{c}, {a} {b}", f"{c} {a} {b}",
                            f"{a[0]}. {c}", f"{a} {b[0]}. {c}"):
                    if len(var) >= 3 and var not in existing and var not in additions:
                        additions[var] = (label, cat)

    entity_map.update(additions)

    # ── Case variants (UPPER / lower) ──
    # Documents often use ALL-CAPS headings: "SPARKASSE" vs "Sparkasse".
    case_adds: Dict[str, Tuple[str, str]] = {}
    for text, (label, cat) in list(entity_map.items()):
        if len(text) < 3 or _is_legal_numbering(text):
            continue
        for variant in (text.upper(), text.lower()):
            if variant != text and variant not in entity_map and variant not in case_adds:
                case_adds[variant] = (label, cat)
    entity_map.update(case_adds)

    # ── Spaced / "gesperrt" variants ──
    # Some documents letter-space names: "S P A R K A S S E" or "M ü l l e r".
    # Generate spaced forms for every single-word entity >= 4 chars.
    spaced_adds: Dict[str, Tuple[str, str]] = {}
    for text, (label, cat) in list(entity_map.items()):
        words = text.split()
        if len(words) != 1:
            continue
        word = words[0]
        if len(word) < 4 or _is_legal_numbering(word):
            continue
        # "Sparkasse" → "S p a r k a s s e"
        spaced = " ".join(word)
        if spaced not in entity_map and spaced not in spaced_adds:
            spaced_adds[spaced] = (label, cat)
        # Also the uppercase spaced form: "S P A R K A S S E"
        spaced_up = " ".join(word.upper())
        if spaced_up not in entity_map and spaced_up != spaced and spaced_up not in spaced_adds:
            spaced_adds[spaced_up] = (label, cat)
    entity_map.update(spaced_adds)


def redact_pdf(
    pdf_path: str,
    output_path: str,
    entity_map: Dict[str, Tuple[str, str]],
    mode: str = "pseudo_vars",
    progress_callback=None,
    api_key: Optional[str] = None,
) -> str:
    """
    Create a redacted copy of *pdf_path* at *output_path*.

    *mode* controls how redactions are rendered (see ``_add_redaction``).
    *entity_map*: ``{original_text: (label, category)}``.

    Returns the output path.
    """
    doc = fitz.open(pdf_path)
    total_pages = len(doc)

    # Pre-scan for repeating images (likely logos / letterheads)
    repeating_xrefs = _find_repeating_image_xrefs(doc)

    # ── Derive implicit PII sub-entities ──
    # If the AI found "Sparkasse Köln-Bonn" as UNTERNEHMEN, we must also
    # catch standalone "Sparkasse Köln-Bonn" fragments and the distinctive
    # core words that clearly identify the same entity.  Similarly for
    # multi-word person names: "Dr. Hans Müller" → also catch "Hans Müller".
    _expand_entity_map(entity_map)

    # Sort entities by length descending so longer matches are processed first.
    # Filter out anything that looks like legal numbering / §§ references.
    sorted_entities = sorted(
        (k for k in entity_map.keys() if not _is_legal_numbering(k)),
        key=len, reverse=True,
    )

    for page_idx, page in enumerate(doc):
        if progress_callback:
            progress_callback(int((page_idx / total_pages) * 100))

        # Collect overlay info for every redaction on this page.
        # After apply_redactions() we re-draw them as shapes to guarantee
        # the black fill is always visible (fixes white-only output on
        # some generated PDFs).
        page_overlays: list = []

        for entity_text in sorted_entities:
            label, category = entity_map[entity_text]

            # Use quads=True for pixel-precise text location, then
            # convert each quad to its enclosing rect with a small
            # vertical pad so the box fully covers the glyphs.
            try:
                quads = page.search_for(entity_text, quads=True)
            except Exception:
                quads = page.search_for(entity_text)

            for q in quads:
                r = q.rect if hasattr(q, "rect") else fitz.Rect(q)
                # Slight vertical padding – covers glyphs fully including
                # descenders/ascenders, without spilling into adjacent lines.
                r = fitz.Rect(r.x0, r.y0 - 1.0, r.x1, r.y1 + 1.0)
                info = _add_redaction(page, r, label, mode, category=category)
                page_overlays.append(info)

        # Detect if this page is a scan (full-page background image).
        # On scans, image/vector-based detection would analyse the scan
        # itself and cause massive false positives → skip those methods.
        page_is_scan = _page_is_scan(page)

        if not page_is_scan:
            # Redact logos / brand images / letterheads in headers/footers
            _redact_logo_images(page, repeating_xrefs, mode)
            # Redact vector drawings in header zone (letterhead graphics)
            _redact_header_zone_drawings(page)

        # Redact signatures, handwriting, ink annotations, etc.
        _detect_and_redact_signatures(page, is_scan=page_is_scan)

        # GPT-5.2 vision: detect signatures, logos, stamps, photos.
        # Only runs on pages likely to contain visual PII (first, last,
        # pages with signature-indicators) to avoid unnecessary API calls.
        if api_key and _page_needs_vision(page, page_idx, total_pages):
            vision_hits = _detect_visuals_with_vision(page, api_key)
            for rect, vis_type in vision_hits:
                margin = 1.0 if vis_type == "paraphe" else _REDACT_MARGIN
                expanded = _safe_expand_rect(rect, page, margin)
                page.add_redact_annot(expanded, text="", fill=REDACT_BG)

        # Collect non-entity redaction rects (signatures, logos) from
        # annotations so we can re-draw them as overlays too.
        entity_keys = {(round(r.x0, 1), round(r.y0, 1),
                        round(r.x1, 1), round(r.y1, 1))
                       for r, _, _, _c in page_overlays}
        annot = page.first_annot
        while annot:
            try:
                if annot.type[0] == 12:  # PDF_ANNOT_REDACT
                    r = fitz.Rect(annot.rect)
                    key = (round(r.x0, 1), round(r.y0, 1),
                           round(r.x1, 1), round(r.y1, 1))
                    if key not in entity_keys:
                        # Non-entity redactions (signatures, logos) → lighter fill
                        page_overlays.append((r, "", 0, ""))
                annot = annot.next
            except Exception:
                break

        # Apply all redactions for this page at once.
        # CRITICAL: On scan pages the entire visible content is one big
        # raster image.  PDF_REDACT_IMAGE_REMOVE would delete it and
        # wipe the whole page.  PDF_REDACT_IMAGE_PIXELS instead blanks
        # only the pixels under each annotation → scan stays intact.
        if page_is_scan:
            page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_PIXELS)
        else:
            page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_REMOVE)

        # Re-draw every redacted area as an elegant filled overlay
        _draw_redaction_overlays(page, page_overlays)

    # -- Strip ALL metadata from output --
    _strip_metadata(doc)

    if progress_callback:
        progress_callback(100)

    doc.save(output_path, garbage=4, deflate=True)
    doc.close()
    return output_path



def get_page_count(pdf_path: str) -> int:
    """Return the number of pages in a PDF."""
    doc = fitz.open(pdf_path)
    count = len(doc)
    doc.close()
    return count
